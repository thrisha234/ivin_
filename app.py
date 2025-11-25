import os
import io
import json
import tempfile
import sqlite3
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file
import numpy as np
import torch
from datetime import datetime, timedelta
import os, sqlite3, json
from flask import Flask, request, render_template, redirect, flash
from werkzeug.utils import secure_filename
import torchaudio
from werkzeug.utils import secure_filename
from pydub import AudioSegment
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from datetime import datetime, timezone
import openai

openai.api_key = os.environ.get("OPENAI_API_KEY")
from reportlab.lib import colors

# Whisper
try:
    import whisper
    WHISPER_AVAILABLE = True
except:
    WHISPER_AVAILABLE = False

# SpeechBrain
try:
    from speechbrain.pretrained import EncoderClassifier
    SPEECHBRAIN_AVAILABLE = True
except:
    SPEECHBRAIN_AVAILABLE = False

# -----------------------------
# PATHS & BASIC SETUP
# -----------------------------
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
DB_DIR = os.path.join(BASE_DIR, "db")

DB_PATH = os.path.join(DB_DIR, "meetings.db")
ENROLL_JSON = os.path.join(DB_DIR, "enrolled.json")

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(DB_DIR, exist_ok=True)

app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = UPLOAD_DIR
app.secret_key = "dev-secret-key"

# -----------------------------
# DATABASE INITIALIZATION (with safe ALTER)
# -----------------------------
def init_db():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("""
    CREATE TABLE IF NOT EXISTS meetings (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        title TEXT,
        date TEXT,
        transcript TEXT,
        minutes_path TEXT,
        docx_path TEXT,
        timeline_path TEXT
    )
    """)
    c.execute("""
    CREATE TABLE IF NOT EXISTS action_items (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        description TEXT,
        raised_in_meeting TEXT,
        meeting_id INTEGER,
        status TEXT DEFAULT 'OPEN',
        created_at TEXT,
        responsible TEXT,
        due_date TEXT,
        challenges TEXT
    )
    """)
    conn.commit()
    conn.close()

# Backwards-safe: if older DB exists without new columns, try to add them
def ensure_columns():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    # get current columns
    c.execute("PRAGMA table_info(meetings)")
    cols = [r[1] for r in c.fetchall()]
    if "docx_path" not in cols:
        try:
            c.execute("ALTER TABLE meetings ADD COLUMN docx_path TEXT")
        except Exception:
            pass
    if "timeline_path" not in cols:
        try:
            c.execute("ALTER TABLE meetings ADD COLUMN timeline_path TEXT")
        except Exception:
            pass
    conn.commit()
    conn.close()

init_db()
ensure_columns()

# -----------------------------
# JSON STORAGE
# -----------------------------
def load_json(path):
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def save_json(path, data):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2)

def load_enrolled():
    return load_json(ENROLL_JSON)

def save_enrolled(data):
    save_json(ENROLL_JSON, data)

# -----------------------------
# MODEL LOADERS
# -----------------------------
_whisper_model = None
_speech_encoder = None

def get_whisper_model(name="small"):
    global _whisper_model
    if _whisper_model is None and WHISPER_AVAILABLE:
        _whisper_model = whisper.load_model(name)
    return _whisper_model

def get_speech_encoder():
    global _speech_encoder
    if _speech_encoder is None and SPEECHBRAIN_AVAILABLE:
        _speech_encoder = EncoderClassifier.from_hparams(
            source="speechbrain/spkrec-ecapa-voxceleb",
            run_opts={"device": "cpu"}
        )
    return _speech_encoder
# Helper function to get Indian time
def current_indian_time():
    utc_now = datetime.utcnow()
    ist_now = utc_now + timedelta(hours=5, minutes=30)
    return ist_now.strftime("%Y-%m-%d"), ist_now.strftime("%H:%M IST")
def compute_embedding_from_file(audio_bytes):
    enc = get_speech_encoder()
    if enc is None:
        return None
    try:
        tf = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
        tf.write(audio_bytes)
        tf.close()
        signal, fs = torchaudio.load(tf.name)
        if signal.shape[0] > 1:
            signal = signal.mean(dim=0, keepdim=True)
        if signal.dim() == 2:
            signal = signal.unsqueeze(0)
        with torch.no_grad():
            embedding = enc.encode_batch(signal).squeeze().cpu().numpy()
        embedding = embedding / (np.linalg.norm(embedding) + 1e-10)
        return embedding
    except Exception as e:
        print("Embedding Error:", e)
        return None
    finally:
        try:
            os.unlink(tf.name)
        except:
            pass

# -----------------------------
# AUDIO UTILITIES
# -----------------------------
def convert_to_wav(input_path):
    audio = AudioSegment.from_file(input_path)
    wav_path = input_path.rsplit(".", 1)[0] + ".wav"
    audio.export(wav_path, format="wav")
    return wav_path

def extract_actions_from_text(text):
    actions = []
    patterns = [
        r'Action[:\-]\s*(.+)',
        r'To do[:\-]\s*(.+)',
        r'action item[:\-]\s*(.+)'
    ]
    lines = re.split(r'[\n\.;]', text)
    for line in lines:
        for p in patterns:
            m = re.search(p, line, re.IGNORECASE)
            if m:
                actions.append(m.group(1).strip())
    return actions

# -----------------------------
# PDF & DOCX generators
# -----------------------------



from docx import Document
from docx.shared import Pt
import os

def generate_minutes_word(meeting_id, subject, summary, attendees, absentees, actions):
    doc_dir = "static/minutes"
    os.makedirs(doc_dir, exist_ok=True)
    doc_path = os.path.join(doc_dir, f"meeting_{meeting_id}.docx")

    doc = Document()
    doc.add_heading("IVIN - Meeting Minutes", 0)
    
    doc.add_paragraph(f"Meeting Subject: {subject}")
    doc.add_paragraph(f"Meeting Date: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph(f"Meeting Time: {datetime.now().strftime('%H:%M IST')}")

    # Attendees
    doc.add_heading("Attendees:", level=1)
    if attendees:
        for person in attendees:
            doc.add_paragraph(f"- {person}")
    else:
        doc.add_paragraph("- Not Provided")

    # Absentees
    doc.add_heading("Absentees:", level=1)
    if absentees:
        for person in absentees:
            doc.add_paragraph(f"- {person}")
    else:
        doc.add_paragraph("- None")

    # Summary
    doc.add_heading("Meeting Summary:", level=1)
    for line in summary.split("\n"):
        doc.add_paragraph(f"- {line}")

    # Action Items
    doc.add_heading("Action Items:", level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "No."
    hdr_cells[1].text = "Owner"
    hdr_cells[2].text = "Task"
    hdr_cells[3].text = "Deadline"

    for idx, action in enumerate(actions, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = action["owner"]
        row_cells[2].text = action["task"]
        row_cells[3].text = action["deadline"]

    doc.save(doc_path)
    return doc_path




# -----------------------------
# MULTI-SPEAKER IDENTIFICATION
# -----------------------------
def identify_speakers_in_audio(audio_bytes,
                               chunk_duration=3.0,
                               overlap=1.5,
                               min_chunk_sec=1.0,
                               similarity_threshold=0.45):

    enc = get_speech_encoder()
    if enc is None:
        print("❌ Speech encoder not loaded")
        return []

    tf = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
    tf.write(audio_bytes)
    tf.close()

    waveform, sr = torchaudio.load(tf.name)

    if waveform.shape[0] > 1:
        waveform = waveform.mean(dim=0, keepdim=True)

    if sr != 16000:
        waveform = torchaudio.transforms.Resample(sr, 16000)(waveform)
        sr = 16000

    total_samples = waveform.shape[1]
    chunk_samples = int(chunk_duration * sr)
    overlap_samples = int(overlap * sr)
    min_samples = int(min_chunk_sec * sr)

    enrolled = load_enrolled()
    timeline = []
    start_idx = 0
    last_speaker = None

    print(f"🎙 Duration: {round(total_samples/sr, 2)}s")

    while start_idx < total_samples:
        end_idx = min(start_idx + chunk_samples, total_samples)
        chunk = waveform[:, start_idx:end_idx]

        if chunk.shape[1] < min_samples:
            chunk = pad_or_trim_waveform(chunk, sr, min_duration=min_chunk_sec)

        chunk = chunk.unsqueeze(0)

        try:
            with torch.no_grad():
                emb_tensor = enc.encode_batch(chunk)
                emb = emb_tensor.squeeze(0).cpu().numpy()
                emb = emb / (np.linalg.norm(emb) + 1e-10)
        except Exception as e:
            print(f"⚠ Encoder error: {e}")
            start_idx += (chunk_samples - overlap_samples)
            continue

        scores = []

        for name, info in enrolled.items():
            e = np.array(info["embedding"])
            e = e / (np.linalg.norm(e) + 1e-10)

            sim = float(np.dot(emb, e))
            scores.append((name, sim))

        if not scores:
            top_name, top_sim = "Unknown", 0.0
        else:
            scores.sort(key=lambda x: x[1], reverse=True)
            top_name, top_sim = scores[0]

        if top_sim > similarity_threshold and (len(scores) == 1 or top_sim - scores[1][1] > 0.05 if len(scores) > 1 else True):
            best_speaker = top_name
        else:
            best_speaker = "Unknown"

        # Force continuity
        if best_speaker == "Unknown" and last_speaker:
            if top_sim > similarity_threshold - 0.1:
                best_speaker = last_speaker

        last_speaker = best_speaker

        timeline.append({
            "start": round(start_idx / sr, 2),
            "end": round(end_idx / sr, 2),
            "speaker": best_speaker,
            "confidence": round(top_sim, 3)
        })

        start_idx += (chunk_samples - overlap_samples)

    try:
        os.unlink(tf.name)
    except:
        pass

    print("✅ Speaker timeline complete.")
    return timeline


# -----------------------------
# Flask API endpoints
# -----------------------------
@app.route("/identify_audio_timeline", methods=["POST"])
def identify_audio_timeline():
    try:
        if "audio" not in request.files:
            return jsonify([])
        file = request.files["audio"]
        audio_bytes = file.read()
        timeline = identify_speakers_in_audio(audio_bytes)
        return jsonify(timeline)
    except Exception as e:
        print("Timeline Identify Error:", e)
        return jsonify([])

@app.route("/")
def index():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT id, description FROM action_items WHERE status='OPEN'")
    actions = c.fetchall()
    conn.close()
    return render_template("index.html", actions=actions)

@app.route("/enroll", methods=["GET", "POST"])
def enroll():
    if request.method == "POST":
        name = request.form.get("name")
        file = request.files.get("audio")
        if not name or not file:
            flash("⚠ Provide both name and audio file")
            return redirect("/enroll")
        filename = secure_filename(file.filename)
        path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
        file.save(path)
        wav_path = convert_to_wav(path)
        try:
            signal, fs = torchaudio.load(wav_path)
            if signal.shape[0] > 1:
                signal = signal.mean(dim=0, keepdim=True)
            if fs != 16000:
                resampler = torchaudio.transforms.Resample(orig_freq=fs, new_freq=16000)
                signal = resampler(signal)
                fs = 16000
            signal = signal.squeeze(0).unsqueeze(0)
            enc = get_speech_encoder()
            if enc is None:
                flash("⚠ Speech encoder model not available!")
                return redirect("/enroll")
            with torch.no_grad():
                emb_tensor = enc.encode_batch(signal)
                emb = emb_tensor.squeeze(0).cpu().numpy()
            emb = emb / (np.linalg.norm(emb) + 1e-10)
            enrolled = load_enrolled()
            enrolled[name] = {
                "embedding": emb.tolist(),
                "file": filename,
                "created_at": datetime.utcnow().isoformat()
            }
            save_enrolled(enrolled)
            flash(f"✅ {name} enrolled at {enrolled[name]['created_at']}")
            return redirect("/enroll")
        except Exception as e:
            print("Failed to compute speaker embedding:", e)
            flash(f"⚠ Failed to compute speaker embedding: {e}")
            return redirect("/enroll")
    enrolled = load_enrolled()
    return render_template("enroll.html", enrolled=enrolled)

@app.route("/delete_speaker", methods=["POST"])
def delete_speaker():
    name = request.form.get("name")
    enrolled = load_enrolled()
    if name in enrolled:
        file_name = enrolled[name].get("file")
        if file_name:
            file_path = os.path.join(app.config["UPLOAD_FOLDER"], file_name)
            if os.path.exists(file_path):
                os.remove(file_path)
        del enrolled[name]
        save_enrolled(enrolled)
        flash(f"🗑 Speaker '{name}' deleted successfully!")
    else:
        flash(f"⚠ Speaker '{name}' not found!")
    return redirect("/enroll")

# -----------------------------
# MEETING UPLOAD
# -----------------------------
@app.route("/upload_meeting", methods=["POST", "GET"])
def upload_meeting():
    if request.method == "POST":
        title = request.form.get("title", "Direction GuideBot - Client Update")
        file = request.files.get("audio")

        if not file:
            flash("⚠ No audio file uploaded")
            return redirect("/upload_meeting")

        filename = secure_filename(file.filename)
        path = os.path.join(UPLOAD_DIR, filename)
        file.save(path)

        wav = convert_to_wav(path)

        # ----------------------------
        # WHISPER TRANSCRIPTION
        # ----------------------------
        whisper_model = get_whisper_model()
        result = whisper_model.transcribe(wav)
        transcript_text = result.get("text", "")

        # ----------------------------
        # CALL OPENAI TO GENERATE FORMATTED MINUTES
        # ----------------------------
        openai.api_key = os.environ.get("OPENAI_API_KEY")
        meeting_date, meeting_time = current_indian_time()

        prompt = f"""
        You are an assistant that generates professional meeting minutes.
        Take the transcript below and generate a readable output in this format:

        IVIN - Meeting Minutes

        Meeting Subject: {title}
        Meeting Date: {meeting_date}
        Meeting Time: {meeting_time}

        Attendees:
        (List all attendees automatically from transcript if possible)

        Meeting Summary:
        - Summarize key discussion points clearly in bullet points

        Action Items:
        No. | Person | Task | Deadline
        (Generate tasks with owner and deadline based on transcript)

        Transcript:
        {transcript_text}
        """

        response = openai.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
        )
        output_text = response.choices[0].message.content

        # ----------------------------
        # PARSE GPT OUTPUT FOR WORD
        # ----------------------------
        def parse_gpt_minutes(output_text):
            lines = output_text.splitlines()
            summary_lines = []
            actions = []
            attendees = []

            in_summary = in_actions = in_attendees = False

            for line in lines:
                line = line.strip()
                if line.startswith("Attendees:"):
                    in_attendees = True
                    in_summary = in_actions = False
                    continue
                elif line.startswith("Meeting Summary"):
                    in_summary = True
                    in_attendees = in_actions = False
                    continue
                elif line.startswith("Action Items"):
                    in_actions = True
                    in_summary = in_attendees = False
                    continue

                if in_attendees and line:
                    attendees.append(line.strip(", "))
                elif in_summary and line:
                    summary_lines.append(line.strip("- ").strip())
                elif in_actions and line and "|" in line:
                    parts = [p.strip() for p in line.split("|")]
                    if len(parts) == 4:
                        actions.append({
                            "task": parts[2],
                            "owner": parts[1],
                            "deadline": parts[3]
                        })

            return attendees, "\n".join(summary_lines), actions

        attendees_list, summary_text, actions_list = parse_gpt_minutes(output_text)

        # ----------------------------
        # SAVE TO DATABASE
        # ----------------------------
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("""
            INSERT INTO meetings (title, date, transcript, docx_path)
            VALUES (?, ?, ?, ?)
        """, (title, datetime.utcnow().isoformat(), transcript_text, ""))
        meeting_id = c.lastrowid
        conn.commit()

        # ----------------------------
        # GENERATE WORD DOCUMENT ONLY
        # ----------------------------
        word_path = generate_minutes_word(
            meeting_id=meeting_id,
            subject=title,
            summary=summary_text,
            attendees=attendees_list,
            absentees=[],
            actions=actions_list
        )

        # ----------------------------
        # UPDATE DB with Word path
        # ----------------------------
        c.execute("""
            UPDATE meetings
            SET docx_path = ?
            WHERE id = ?
        """, (word_path, meeting_id))
        conn.commit()
        conn.close()

        return render_template(
            "meeting_processed.html",
            title=title,
            transcript=transcript_text,
            word_path=word_path,
            meeting_id=meeting_id
        )

    return render_template("upload_meeting.html")


def transcript_to_line_paragraph(labeled_transcript):
    output = []
    current_speaker = None

    for entry in labeled_transcript:
        speaker = entry["speaker"]
        text = entry["text"].strip()

        # New speaker → start a new paragraph
        if speaker != current_speaker:
            output.append("")  # blank line between speakers
            output.append(f"{speaker}: {text}")
            current_speaker = speaker
        else:
            # Same speaker → new line but same paragraph idea
            output.append(text)

    return "\n".join(output).strip()

def pad_or_trim_waveform(waveform, sr, min_duration=1.0):
    """
    Pads waveform to a minimum duration to prevent SpeechBrain padding errors.
    """
    min_samples = int(sr * min_duration)
    if waveform.shape[1] < min_samples:
        pad_size = min_samples - waveform.shape[1]
        waveform = torch.nn.functional.pad(waveform, (0, pad_size))
    return waveform

# -----------------------------
# DOWNLOAD ROUTES
# -----------------------------


@app.route("/download_minutes_docx/<int:meeting_id>")
def download_minutes_docx(meeting_id):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT docx_path FROM meetings WHERE id=?", (meeting_id,))
    row = c.fetchone()
    conn.close()
    if not row or not row[0] or not os.path.exists(row[0]):
        flash("DOCX Not Found")
        return redirect("/")
    return send_file(row[0], as_attachment=True)



# -----------------------------
# RUN SERVER
# -----------------------------
if __name__ == "__main__":
    app.run(debug=True, port=5000)
